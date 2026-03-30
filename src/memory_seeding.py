from __future__ import annotations

import argparse
import os
import sys
import time
from collections import Counter
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Sequence

import httpx

from src.llm_client import LLMClient
from src.memory_store import MemoryConfig, MemoryStore, _utc_now_iso, tokenize


@dataclass(frozen=True)
class SeedStats:
    chunks_processed: int = 0
    translated_chunks: int = 0
    concepts_seeded: int = 0
    bases_seeded: int = 0
    anchored_edges_written: int = 0


def _iter_docx_chunks(path: Path) -> list[str]:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "python-docx is required for .docx input. Install with: pip install python-docx"
        ) from exc

    doc = Document(str(path))
    chunks: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [((cell.text or "").strip()) for cell in row.cells]
            row_text = " | ".join([c for c in cells if c])
            if row_text:
                chunks.append(row_text)

    return chunks


def _iter_path_chunks(path: Path) -> list[str]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _iter_docx_chunks(path)
    raw = path.read_text(encoding="utf-8")
    return [line.strip() for line in raw.splitlines() if line.strip()]


def _translate_chunk(
    chunk: str,
    llm: LLMClient,
    *,
    max_retries: int = 3,
) -> str:
    prompt = (
        "Translate the following text to concise natural English.\n"
        "Preserve meaning and details, and return only the translated text.\n\n"
        f"{chunk}"
    )
    last: BaseException | None = None
    for attempt in range(max(1, max_retries)):
        try:
            result = llm.generate(prompt=prompt)
            translated = (result.get("completion") or "").strip()
            return translated or chunk
        except (httpx.ReadTimeout, httpx.ConnectTimeout, httpx.ConnectError) as exc:
            last = exc
            if attempt >= max(1, max_retries) - 1:
                raise
            time.sleep(min(60.0, 2.0**attempt))
    raise last  # pragma: no cover


def _select_seed_tokens(
    text: str,
    *,
    min_token_len: int,
    max_tokens_per_chunk: int,
) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for tok in tokenize(text):
        t = tok.strip().lower()
        if len(t) < min_token_len:
            continue
        if t in seen:
            continue
        seen.add(t)
        out.append(t)
        if len(out) >= max_tokens_per_chunk:
            break
    return out


def _build_store(memory_db_dir: str | None) -> MemoryStore:
    if not memory_db_dir:
        return MemoryStore()
    db_dir = Path(memory_db_dir).resolve()
    cfg = MemoryConfig(db_path=(db_dir / "memory.sqlite3"))
    return MemoryStore(cfg=cfg)


def seed_concepts_from_paths(
    *,
    input_paths: Sequence[Path],
    store: MemoryStore | None,
    translate: bool,
    llm: LLMClient | None,
    min_token_len: int = 3,
    max_tokens_per_chunk: int = 30,
    translate_retries: int = 3,
    dry_run: bool = False,
) -> tuple[Counter[str], SeedStats]:
    token_freq: Counter[str] = Counter()
    chunks_processed = 0
    translated_chunks = 0
    seeded_names: set[str] = set()

    chunk_seq = 0
    for path in input_paths:
        chunks = _iter_path_chunks(path)
        for chunk in chunks:
            if not chunk:
                continue
            chunks_processed += 1
            chunk_seq += 1
            print(
                f"[memory_seeding] chunk {chunk_seq}: {path.name} ({len(chunk)} chars) "
                f"translate={translate}",
                file=sys.stderr,
                flush=True,
            )
            text = chunk
            if translate:
                if llm is None:
                    raise RuntimeError("Translation requested but no LLM client is configured.")
                text = _translate_chunk(chunk, llm, max_retries=translate_retries)
                translated_chunks += 1

            selected = _select_seed_tokens(
                text,
                min_token_len=min_token_len,
                max_tokens_per_chunk=max_tokens_per_chunk,
            )
            token_freq.update(selected)

            if dry_run:
                continue

            assert store is not None
            snippet = text[:220].replace("\n", " ").strip()
            for tok in selected:
                seeded_names.add(tok)
                summary = (
                    f"Seeded from source text. Example context: {snippet}"
                    if snippet
                    else ""
                )
                store.upsert_node(name=tok, type_="concept", summary=summary)

    stats = SeedStats(
        chunks_processed=chunks_processed,
        translated_chunks=translated_chunks,
        concepts_seeded=len(seeded_names),
    )
    return token_freq, stats


def seed_base_nodes(
    *,
    base_names: Sequence[str],
    store: MemoryStore | None,
    dry_run: bool,
) -> int:
    names = sorted({(b or "").strip().lower() for b in base_names if (b or "").strip()})
    if dry_run:
        return len(names)
    assert store is not None
    for name in names:
        store.upsert_node(name=name, type_="base", summary=f"Seeded base pillar: {name}")
    return len(names)


def cluster_and_anchor(
    *,
    store: MemoryStore | None,
    token_freq: Counter[str],
    similarity_threshold: float = 0.7,
    relation_type: str = "anchored_to",
    write_embeddings: bool = False,
    dry_run: bool = False,
) -> tuple[int, int]:
    if dry_run:
        return 0, 0
    assert store is not None

    try:
        import numpy as np  # type: ignore
        from sentence_transformers import SentenceTransformer  # type: ignore
    except ImportError as exc:  # pragma: no cover - import guard
        raise RuntimeError(
            "Clustering requires sentence-transformers and numpy. "
            "Install with: pip install -r requirements-seeding.txt"
        ) from exc

    rows = [r for r in store.fetch_all_nodes() if str(r["type"] or "").lower() == "concept"]
    if not rows:
        return 0, 0

    concept_ids = [int(r["id"]) for r in rows]
    concept_names = [str(r["name"] or "").strip().lower() for r in rows]
    model = SentenceTransformer("all-MiniLM-L6-v2")
    vectors = model.encode(concept_names)
    vecs = np.asarray(vectors)

    if write_embeddings:
        for idx, name in enumerate(concept_names):
            emb_bytes = np.asarray(vecs[idx], dtype=np.float32).tobytes()
            store.upsert_node(name=name, type_="concept", embedding_blob=emb_bytes)

    clusters: list[list[int]] = []
    used: set[int] = set()
    for i in range(len(vecs)):
        if i in used:
            continue
        cluster = [i]
        used.add(i)
        for j in range(i + 1, len(vecs)):
            if j in used:
                continue
            a = vecs[i]
            b = vecs[j]
            denom = float(np.linalg.norm(a) * np.linalg.norm(b))
            if denom <= 0.0:
                continue
            sim = float(np.dot(a, b) / denom)
            if sim >= similarity_threshold:
                cluster.append(j)
                used.add(j)
        clusters.append(cluster)

    ts = _utc_now_iso()
    relation = (relation_type or "anchored_to").strip().lower()
    base_count = 0
    edge_count = 0

    for cluster in clusters:
        cluster_names = [concept_names[i] for i in cluster]
        candidates = [n for n in cluster_names if len(n) >= 3]
        if not candidates:
            continue
        # Prefer frequent semantic token; tie-break by shorter and then lexical order.
        candidates.sort(key=lambda n: (-int(token_freq.get(n, 0)), len(n), n))
        base_name = candidates[0]
        base_id = store.upsert_node(
            name=base_name,
            type_="base",
            summary=f"Auto-clustered base pillar derived from concept neighborhood: {base_name}",
        )
        base_count += 1

        for i in cluster:
            concept_id = concept_ids[i]
            if concept_id == base_id:
                continue
            store.upsert_edge(
                src_id=concept_id,
                dst_id=base_id,
                relation_type=relation,
                weight_delta=1.0,
                ts=ts,
            )
            edge_count += 1

    return base_count, edge_count


def _parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        prog="python -m src.memory_seeding",
        description="Seed memory.sqlite3 from DOCX/text sources for the agent memory graph.",
    )
    parser.add_argument(
        "--inputs",
        nargs="+",
        required=True,
        help="One or more input files (.docx, .txt, .md).",
    )
    parser.add_argument(
        "--memory-db-dir",
        default=None,
        help="Override MEMORY_DB_DIR for this command. Database is <dir>/memory.sqlite3.",
    )
    parser.add_argument(
        "--translate",
        action="store_true",
        help="Translate each source chunk to English before token seeding.",
    )
    parser.add_argument(
        "--translate-model",
        default="qwen3:4b",
        help="Model passed to LLMClient when --translate is enabled.",
    )
    parser.add_argument(
        "--translate-base-url",
        default=None,
        help="Optional Ollama base URL override for translation.",
    )
    parser.add_argument(
        "--translate-timeout",
        type=float,
        default=None,
        help=(
            "HTTP timeout (seconds) for each Ollama request during --translate. "
            "Default: OLLAMA_SEED_TIMEOUT env or 600 (longer than chat; avoids httpx.ReadTimeout on slow loads)."
        ),
    )
    parser.add_argument(
        "--translate-retries",
        type=int,
        default=3,
        help="Retries on connection/read timeout when calling Ollama for --translate.",
    )
    parser.add_argument(
        "--min-token-len",
        type=int,
        default=3,
        help="Minimum token length kept for seeding.",
    )
    parser.add_argument(
        "--max-tokens-per-chunk",
        type=int,
        default=30,
        help="Maximum unique tokens persisted per source chunk.",
    )
    parser.add_argument(
        "--base-name",
        action="append",
        default=[],
        help="Manual base pillar name; can be repeated.",
    )
    parser.add_argument(
        "--cluster-bases",
        action="store_true",
        help="Enable embedding-based concept clustering to create base nodes and anchored_to edges.",
    )
    parser.add_argument(
        "--cluster-threshold",
        type=float,
        default=0.7,
        help="Cosine similarity threshold for cluster grouping.",
    )
    parser.add_argument(
        "--write-embeddings",
        action="store_true",
        help="Persist concept embeddings into nodes.embedding_blob during clustering.",
    )
    parser.add_argument(
        "--relation-type",
        default="anchored_to",
        help="Relation type used for concept-to-base links.",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse and tokenize inputs only; do not write to SQLite.",
    )
    return parser.parse_args(argv)


def _validate_paths(raw_paths: Sequence[str]) -> list[Path]:
    paths: list[Path] = []
    for p in raw_paths:
        path = Path(p).resolve()
        if not path.exists():
            raise FileNotFoundError(f"Input path not found: {path}")
        if not path.is_file():
            raise RuntimeError(f"Input path is not a file: {path}")
        paths.append(path)
    return paths


def run(argv: Sequence[str] | None = None) -> SeedStats:
    args = _parse_args(argv)
    input_paths = _validate_paths(args.inputs)

    store: MemoryStore | None = None
    if not args.dry_run:
        if args.memory_db_dir:
            os.environ["MEMORY_DB_DIR"] = str(Path(args.memory_db_dir).resolve())
        store = _build_store(args.memory_db_dir)

    llm: LLMClient | None = None
    if args.translate:
        t_out = args.translate_timeout
        if t_out is None:
            t_out = float(os.getenv("OLLAMA_SEED_TIMEOUT", "600"))
        llm = LLMClient(
            base_url=args.translate_base_url,
            model=args.translate_model,
            timeout=float(t_out),
        )

    token_freq, stats = seed_concepts_from_paths(
        input_paths=input_paths,
        store=store,
        translate=bool(args.translate),
        llm=llm,
        min_token_len=max(1, int(args.min_token_len)),
        max_tokens_per_chunk=max(1, int(args.max_tokens_per_chunk)),
        translate_retries=max(1, int(args.translate_retries)),
        dry_run=bool(args.dry_run),
    )

    base_count = seed_base_nodes(
        base_names=args.base_name,
        store=store,
        dry_run=bool(args.dry_run),
    )

    clustered_base_count = 0
    anchored_count = 0
    if args.cluster_bases:
        clustered_base_count, anchored_count = cluster_and_anchor(
            store=store,
            token_freq=token_freq,
            similarity_threshold=float(args.cluster_threshold),
            relation_type=args.relation_type,
            write_embeddings=bool(args.write_embeddings),
            dry_run=bool(args.dry_run),
        )

    final_stats = SeedStats(
        chunks_processed=stats.chunks_processed,
        translated_chunks=stats.translated_chunks,
        concepts_seeded=stats.concepts_seeded,
        bases_seeded=int(base_count) + int(clustered_base_count),
        anchored_edges_written=anchored_count,
    )
    return final_stats


def _print_stats(stats: SeedStats) -> None:
    print(
        "Seeding complete | "
        f"chunks={stats.chunks_processed} "
        f"translated={stats.translated_chunks} "
        f"concepts={stats.concepts_seeded} "
        f"bases={stats.bases_seeded} "
        f"anchored_edges={stats.anchored_edges_written}"
    )


if __name__ == "__main__":
    out = run()
    _print_stats(out)
